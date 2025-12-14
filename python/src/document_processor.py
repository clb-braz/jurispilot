"""
JurisPilot - Processador de Documentos
Extrai texto, metadados e identifica tipos de documentos jurídicos
"""

import os
import json
from typing import Dict, Optional, List
from pathlib import Path
from datetime import datetime
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from loguru import logger
import dateparser


class DocumentProcessor:
    """Processa documentos jurídicos extraindo texto, metadados e identificando tipos"""
    
    # Tipos de documentos jurídicos conhecidos
    DOCUMENT_TYPES = {
        'cpf': ['cpf', 'cadastro', 'pessoa física'],
        'cnpj': ['cnpj', 'cadastro nacional', 'pessoa jurídica'],
        'rg': ['rg', 'registro geral', 'identidade'],
        'contrato': ['contrato', 'termo', 'acordo'],
        'holerite': ['holerite', 'contracheque', 'recibo de pagamento'],
        'extrato_bancario': ['extrato', 'banco', 'movimentação'],
        'nota_fiscal': ['nota fiscal', 'nf', 'nfe'],
        'boleto': ['boleto', 'cobrança', 'pagamento'],
        'certidao': ['certidão', 'certidão de nascimento', 'certidão de casamento'],
        'irpf': ['irpf', 'imposto de renda', 'declaração'],
        'carteira_trabalho': ['ctps', 'carteira de trabalho'],
        'email': ['email', 'e-mail', 'mensagem'],
        'protocolo': ['protocolo', 'número de protocolo'],
        'comprovante': ['comprovante', 'recibo', 'comprovante de pagamento']
    }
    
    def __init__(self, tesseract_path: Optional[str] = None):
        """Inicializa o processador de documentos"""
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        else:
            # Detecção automática de Tesseract
            tesseract_path = self._detect_tesseract()
            if tesseract_path:
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
        logger.info("DocumentProcessor inicializado")
    
    @staticmethod
    def _detect_tesseract() -> Optional[str]:
        """Detecta automaticamente o caminho do Tesseract em diferentes sistemas operacionais"""
        import platform
        import shutil
        
        # Verifica se está no PATH
        tesseract_path = shutil.which("tesseract")
        if tesseract_path:
            return tesseract_path
        
        # Paths específicos por OS
        system = platform.system()
        
        if system == "Windows":
            # Windows - locais comuns
            possible_paths = [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                os.path.expanduser(r"~\AppData\Local\Tesseract-OCR\tesseract.exe"),
            ]
        elif system == "Darwin":  # macOS
            possible_paths = [
                "/usr/local/bin/tesseract",
                "/opt/homebrew/bin/tesseract",
                "/usr/bin/tesseract",
            ]
        else:  # Linux
            possible_paths = [
                "/usr/bin/tesseract",
                "/usr/local/bin/tesseract",
            ]
        
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        return None
    
    def process_file(self, file_path: str) -> Dict:
        """
        Processa um arquivo e retorna informações extraídas
        
        Args:
            file_path: Caminho do arquivo a processar
            
        Returns:
            Dict com texto, metadados e tipo de documento
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        file_name = Path(file_path).name
        
        logger.info(f"Processando arquivo: {file_name}")
        
        result = {
            'nome_arquivo': file_name,
            'caminho_arquivo': file_path,
            'tamanho_arquivo': os.path.getsize(file_path),
            'mime_type': self._get_mime_type(file_ext),
            'tipo_documento': None,
            'texto_extraido': '',
            'metadados': {},
            'data_documento': None,
            'valores_encontrados': [],
            'partes_envolvidas': []
        }
        
        # Extrai texto baseado na extensão
        if file_ext == '.pdf':
            result.update(self._process_pdf(file_path))
        elif file_ext in ['.doc', '.docx']:
            result.update(self._process_docx(file_path))
        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
            result.update(self._process_image(file_path))
        else:
            logger.warning(f"Tipo de arquivo não suportado: {file_ext}")
            result['texto_extraido'] = f"Tipo de arquivo {file_ext} não suportado para extração de texto"
        
        # Identifica tipo de documento
        result['tipo_documento'] = self._identify_document_type(result['texto_extraido'])
        
        # Extrai metadados adicionais
        result['metadados'] = self._extract_metadata(result['texto_extraido'])
        result['data_documento'] = self._extract_date(result['texto_extraido'])
        result['valores_encontrados'] = self._extract_values(result['texto_extraido'])
        
        logger.info(f"Processamento concluído: {file_name} - Tipo: {result['tipo_documento']}")
        
        return result
    
    def _process_pdf(self, file_path: str) -> Dict:
        """Processa arquivo PDF"""
        text = ""
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['num_paginas'] = len(pdf_reader.pages)
                
                # Extrai texto de todas as páginas
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    text += f"\n--- Página {page_num} ---\n{page_text}"
                
                # Tenta extrair metadados do PDF
                if pdf_reader.metadata:
                    metadata['pdf_metadata'] = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', '')
                    }
        except Exception as e:
            logger.error(f"Erro ao processar PDF: {e}")
            # Tenta OCR se a extração de texto falhar
            try:
                images = convert_from_path(file_path)
                text = ""
                for img in images:
                    text += pytesseract.image_to_string(img, lang='por') + "\n"
            except Exception as ocr_error:
                logger.error(f"Erro no OCR: {ocr_error}")
                text = f"Erro ao processar PDF: {str(e)}"
        
        return {'texto_extraido': text.strip(), 'metadados': metadata}
    
    def _process_docx(self, file_path: str) -> Dict:
        """Processa arquivo Word"""
        text = ""
        metadata = {}
        
        try:
            doc = Document(file_path)
            
            # Extrai texto de todos os parágrafos
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extrai texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Metadados do documento
            if doc.core_properties:
                metadata['docx_metadata'] = {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
                }
        except Exception as e:
            logger.error(f"Erro ao processar DOCX: {e}")
            text = f"Erro ao processar documento Word: {str(e)}"
        
        return {'texto_extraido': text.strip(), 'metadados': metadata}
    
    def _process_image(self, file_path: str) -> Dict:
        """Processa imagem usando OCR"""
        text = ""
        metadata = {}
        
        try:
            image = Image.open(file_path)
            metadata['image_size'] = image.size
            metadata['image_format'] = image.format
            
            # OCR
            text = pytesseract.image_to_string(image, lang='por')
        except Exception as e:
            logger.error(f"Erro ao processar imagem: {e}")
            text = f"Erro ao processar imagem: {str(e)}"
        
        return {'texto_extraido': text.strip(), 'metadados': metadata}
    
    def _get_mime_type(self, file_ext: str) -> str:
        """Retorna MIME type baseado na extensão"""
        mime_types = {
            '.pdf': 'application/pdf',
            '.doc': 'application/msword',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.bmp': 'image/bmp',
            '.tiff': 'image/tiff'
        }
        return mime_types.get(file_ext, 'application/octet-stream')
    
    def _identify_document_type(self, text: str) -> Optional[str]:
        """Identifica o tipo de documento baseado no texto"""
        text_lower = text.lower()
        
        for doc_type, keywords in self.DOCUMENT_TYPES.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return doc_type
        
        return 'documento_generico'
    
    def _extract_metadata(self, text: str) -> Dict:
        """Extrai metadados do texto"""
        metadata = {
            'num_palavras': len(text.split()),
            'num_linhas': len(text.split('\n')),
            'tem_cpf': self._has_cpf(text),
            'tem_cnpj': self._has_cnpj(text),
            'tem_email': self._has_email(text),
            'tem_telefone': self._has_phone(text)
        }
        return metadata
    
    def _extract_date(self, text: str) -> Optional[str]:
        """Extrai data do documento"""
        # Usa dateparser para encontrar datas no texto
        try:
            # Procura por padrões de data comuns
            date_patterns = [
                r'\d{2}/\d{2}/\d{4}',
                r'\d{2}-\d{2}-\d{4}',
                r'\d{4}-\d{2}-\d{2}',
                r'\d{1,2}\s+de\s+\w+\s+de\s+\d{4}'
            ]
            
            import re
            for pattern in date_patterns:
                matches = re.findall(pattern, text)
                if matches:
                    parsed_date = dateparser.parse(matches[0], languages=['pt'])
                    if parsed_date:
                        return parsed_date.strftime('%Y-%m-%d')
            
            # Tenta parsear qualquer data no texto
            parsed = dateparser.parse(text, languages=['pt'])
            if parsed:
                return parsed.strftime('%Y-%m-%d')
        except Exception as e:
            logger.debug(f"Erro ao extrair data: {e}")
        
        return None
    
    def _extract_values(self, text: str) -> List[float]:
        """Extrai valores monetários do texto"""
        import re
        values = []
        
        # Padrões para valores monetários (R$ 1.234,56 ou 1234.56)
        patterns = [
            r'R\$\s*(\d{1,3}(?:\.\d{3})*(?:,\d{2})?)',
            r'(\d{1,3}(?:\.\d{3})*(?:,\d{2})?)\s*reais',
            r'valor[:\s]+R\$\s*(\d{1,3}(?:\.\d{3})*(?:,\d{2})?)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    # Converte formato brasileiro para float
                    value_str = match.replace('.', '').replace(',', '.')
                    value = float(value_str)
                    values.append(value)
                except:
                    pass
        
        return list(set(values))  # Remove duplicatas
    
    def _has_cpf(self, text: str) -> bool:
        """Verifica se o texto contém CPF"""
        import re
        cpf_pattern = r'\d{3}\.?\d{3}\.?\d{3}-?\d{2}'
        return bool(re.search(cpf_pattern, text))
    
    def _has_cnpj(self, text: str) -> bool:
        """Verifica se o texto contém CNPJ"""
        import re
        cnpj_pattern = r'\d{2}\.?\d{3}\.?\d{3}/?\d{4}-?\d{2}'
        return bool(re.search(cnpj_pattern, text))
    
    def _has_email(self, text: str) -> bool:
        """Verifica se o texto contém email"""
        import re
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        return bool(re.search(email_pattern, text))
    
    def _has_phone(self, text: str) -> bool:
        """Verifica se o texto contém telefone"""
        import re
        phone_pattern = r'\(?\d{2}\)?\s?\d{4,5}-?\d{4}'
        return bool(re.search(phone_pattern, text))


if __name__ == "__main__":
    # Exemplo de uso
    processor = DocumentProcessor()
    
    # Teste com arquivo (se fornecido)
    import sys
    if len(sys.argv) > 1:
        result = processor.process_file(sys.argv[1])
        print(json.dumps(result, indent=2, ensure_ascii=False))

